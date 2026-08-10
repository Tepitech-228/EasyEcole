# -*- coding: utf-8 -*-
"""Génère un document Word (.docx) qui découpe le projet EasyEcole en processus métier."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BLEU = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x59, 0x59, 0x59)
VERT = RGBColor(0x2E, 0x7D, 0x32)

PÔLES = [
    {
        "titre": "PÔLE PÉDAGOGIQUE",
        "processus": [
            {
                "nom": "1. Inscription",
                "etapes": [
                    "Création de session d'inscription (année académique + niveau)",
                    "Demande d'inscription par l'apprenant (choix de session)",
                    "Choix de parcours",
                    "Upload des pièces justificatives",
                    "Soumission de la pré-inscription",
                    "Validation par le comité d'orientation",
                    "Choix des cours (obligatoires + optionnels)",
                    "Paiement des frais d'inscription (espèce ou en ligne)",
                    "Validation finale de l'inscription",
                    "Création du dossier étudiant (matricule + QR code)",
                    "Génération de la carte étudiant",
                    "Création automatique du compte parent (si email fourni)",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "2. Gestion des cours",
                "etapes": [
                    "Planification des cours par l'institution",
                    "Affectation des enseignants aux cours",
                    "Création des chapitres et ressources pédagogiques",
                    "Planification des séances (créneaux + salles)",
                    "Signature des feuilles de présence (QR code étudiant)",
                    "Remplissage du cahier de texte",
                    "Suivi des présences / absences",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "3. Évaluation & Gestion des notes",
                "etapes": [
                    "Configuration des MCC (Modalités de Contrôle des Connaissances)",
                    "Configuration des échelles de notation",
                    "Saisie des notes par les enseignants",
                    "Calcul des moyennes (UE, semestre, générale)",
                    "Affichage notes < 10 en rouge",
                    "Délibérations du jury",
                    "Génération des bulletins et relevés de notes",
                    "Gestion des rattrapages",
                    "Équivalences et dispenses",
                    "Gestion des dettes académiques",
                    "Clôture de semestre : gel des notes, passation (en cours)",
                ],
                "etat": "EN COURS (clôture semestre)",
            },
            {
                "nom": "4. Stages",
                "etapes": [
                    "Dépôt d'offres de stage par les entreprises",
                    "Candidature des étudiants",
                    "Convention de stage",
                    "Suivi et rapport de stage",
                    "Évaluation du stage",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "5. Scolarité",
                "etapes": [
                    "Demandes de documents (diplômes, attestations)",
                    "Traitement des réclamations étudiantes",
                    "Registres académiques (filière/niveau, génération auto depuis délibération)",
                    "Discipline et sanctions",
                    "Conseils de classe et décisions de passage",
                    "Réorientation",
                    "VAE (Validation des Acquis de l'Expérience)",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "6. Paiements & échéances de scolarité",
                "etapes": [
                    "Configuration des frais de parcours",
                    "Génération des bordereaux de paiement",
                    "Encaissement (espèce / en ligne) et quittances",
                    "Vérification du paiement au scan QR (bip aigu / double-bip grave)",
                    "Alertes d'échéances (cron quotidien + notifications + widget dashboard)",
                    "Relances et suivi des dettes",
                    "Intégration comptable automatique des paiements",
                ],
                "etat": "LIVRÉ",
            },
        ],
    },
    {
        "titre": "PÔLE FINANCIER / ERP COMPTABLE",
        "processus": [
            {
                "nom": "1. Comptabilité générale",
                "etapes": [
                    "Plan comptable paramétrable (classes 1 à 9, SYSCOHADA)",
                    "Saisie des écritures comptables",
                    "Journal général, grand livre, balance",
                    "Intégration automatique des paiements d'inscription",
                    "Intégration automatique des bulletins de paie",
                    "Bilan et compte de résultat (format simplifié existant)",
                ],
                "etat": "LIVRÉ (partiel)",
            },
            {
                "nom": "2. États financiers SYSCOHADA (à planifier)",
                "etapes": [
                    "S0 : Référentiel SYSCOHADA, postes de synthèse",
                    "S1 : Bilan + compte de résultat au format officiel OHADA",
                    "S2 : TAFIRE (tableau des ressources et emplois)",
                    "S3 : ETATC, notes annexes essentielles",
                    "S4 : Assemblage du dossier PDF complet (60+ pages)",
                ],
                "etat": "A PLANIFIER",
            },
            {
                "nom": "3. Comptabilité complémentaire (à planifier)",
                "etapes": [
                    "F1 : Clôture d'exercice automatique (résultat, report à nouveau)",
                    "F2 : Flux de trésorerie, SIG et KPIs financiers",
                    "F3 : Connexions modules (stocks, marchés, stages) + TVA",
                ],
                "etat": "A PLANIFIER",
            },
            {
                "nom": "4. Achats",
                "etapes": [
                    "Demande d'achat (besoin)",
                    "Validation hiérarchique",
                    "Demande de prix / devis",
                    "Commande fournisseur",
                    "Réception de la commande",
                    "Facturation",
                    "Suivi budget vs réel",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "5. Stocks",
                "etapes": [
                    "Gestion des articles (cycle de vie)",
                    "Mouvements de stock (entrée / sortie)",
                    "Gestion des fournisseurs",
                    "Inventaires",
                    "Corrections de stock",
                    "Transferts entre sites",
                    "Rebuts",
                    "Affectation des articles aux salles de classe",
                    "Reporting stock",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "6. Immobilisations",
                "etapes": [
                    "Inventaire des biens",
                    "Suivi des immobilisations",
                    "Amortissements (lien comptabilité partiel)",
                ],
                "etat": "LIVRÉ (partiel)",
            },
        ],
    },
    {
        "titre": "PÔLE RESSOURCES HUMAINES",
        "processus": [
            {
                "nom": "1. Gestion des employés & contrats",
                "etapes": [
                    "Fiches employés",
                    "Catégories professionnelles",
                    "Contrats enseignants (création, activation, résiliation + archivage GED)",
                    "Grilles salariales",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "2. Recrutement",
                "etapes": [
                    "Publication d'offres d'emploi",
                    "Dépôt de candidatures",
                    "Présélection et entretiens (workflow à compléter)",
                    "Décision et contrat",
                ],
                "etat": "EN COURS (workflow à compléter)",
            },
            {
                "nom": "3. Pointage & présences",
                "etapes": [
                    "Terminal de pointage (arrivée / départ)",
                    "Gestion des shifts",
                    "Suivi des absences",
                    "Agrégation des heures pour la paie (source des HAO)",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "4. Paie & bulletins",
                "etapes": [
                    "Bulletin de paie : brouillon → validé → versé",
                    "Heures supplémentaires",
                    "Prêts / remboursements",
                    "Écritures comptables automatiques à la génération",
                    "Indemnités prestataires",
                    "R5 : charges sociales Togo (CNSS, IR) — à planifier",
                ],
                "etat": "EN COURS (complément Togo à planifier)",
            },
            {
                "nom": "5. Congés & absences (à compléter)",
                "etapes": [
                    "Demande de congé (workflow validation hiérarchique manquant)",
                    "Suivi des soldes de congés",
                    "Reporting des absences",
                ],
                "etat": "A COMPLÉTER",
            },
            {
                "nom": "6. Programme RH / HAO (à planifier)",
                "etapes": [
                    "R1 : Référentiel RH paramétrable (statuts, quotas, taux horaires HAO, charges)",
                    "R2 : Extraction des heures depuis les pointages",
                    "R3 : Calcul automatique des salaires prestataires + écritures auto",
                    "R4 : Suivi HAO annuel (quota consommé, alertes dépassement)",
                    "R5 : Complétude paie Togo (CNSS, IR, congés payés)",
                ],
                "etat": "A PLANIFIER",
            },
        ],
    },
    {
        "titre": "PÔLE COMMUNICATION",
        "processus": [
            {
                "nom": "1. Messagerie interne",
                "etapes": [
                    "Messages entre utilisateurs",
                    "Chats / salons en temps réel (socket + SSE)",
                    "Annonces et suggestions",
                    "Notifications",
                ],
                "etat": "LIVRÉ",
            },
        ],
    },
    {
        "titre": "PÔLE GED — ARCHIVAGE NUMÉRIQUE",
        "processus": [
            {
                "nom": "1. Archivage des documents",
                "etapes": [
                    "Stockage : public/dossiers/{annee}/{parcours}/{classe}/{niveau}/{matricule}/",
                    "Copie GED + références en base",
                    "Courrier entrant / sortant",
                    "Règles de conservation",
                    "Dossiers arborescents (navigation par compteurs)",
                ],
                "etat": "LIVRÉ",
            },
        ],
    },
    {
        "titre": "PÔLE E-LEARNING",
        "processus": [
            {
                "nom": "1. Cours en ligne",
                "etapes": [
                    "Création et gestion des cours en ligne",
                    "Contenu : vidéos, PDFs, supports de cours",
                    "Chat par cours",
                    "Diffusion en direct (player)",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "2. Devoirs & Quiz",
                "etapes": [
                    "Quiz (QCM, questions ouvertes)",
                    "Devoirs à rendre et soumissions",
                    "Correction et notes",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "3. Progression & Certificats",
                "etapes": [
                    "Suivi de la progression apprenant",
                    "Certificats de réussite",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "4. Arborescence e-learning (à planifier)",
                "etapes": [
                    "EL1 : Rattachement hiérarchique (Année → Parcours → Filière → Niveau)",
                    "EL2 : Catalogue arborescent avec compteurs",
                    "EL3 : Arborescence par entité (devoirs, quiz, supports, certificats)",
                ],
                "etat": "A PLANIFIER",
            },
        ],
    },
    {
        "titre": "PÔLE ESPACE PARENTS",
        "processus": [
            {
                "nom": "1. Suivi de l'étudiant par le parent",
                "etapes": [
                    "Compte parent automatique à l'inscription",
                    "Suivi des notes",
                    "Suivi des absences",
                    "Suivi des paiements / échéances",
                    "Alertes d'échéances sur le dashboard",
                ],
                "etat": "LIVRÉ",
            },
        ],
    },
    {
        "titre": "PÔLE ADMINISTRATION & SYSTÈME",
        "processus": [
            {
                "nom": "1. Gestion des utilisateurs & rôles",
                "etapes": [
                    "Utilisateurs et authentification",
                    "Rôles et permissions",
                    "Configuration de l'établissement",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "2. QR codes & scan",
                "etapes": [
                    "Génération QR codes (étudiant, carte, documents)",
                    "Vérification paiement au scan",
                    "Vérification publique des documents (anti-fraude HMAC SHA-256)",
                ],
                "etat": "LIVRÉ",
            },
            {
                "nom": "3. Génération de documents (DocGen)",
                "etapes": [
                    "Configuration des types de documents",
                    "Création des templates HTML (variables, boucles, conditions)",
                    "Génération PDF (relevés, attestations, diplômes, PV, décisions)",
                    "Signature (enseignant → direction) et cachets électroniques",
                    "QR code anti-fraude",
                    "Vérification publique /verification/document/:matricule/:reference",
                ],
                "etat": "LIVRÉ",
            },
        ],
    },
    {
        "titre": "PROGRAMMES TRANSVERSES",
        "processus": [
            {
                "nom": "1. Arborescence généralisée (à planifier)",
                "etapes": [
                    "Étudiants : Année → Filière → Classe → Étudiant",
                    "Présences : Classe → Cours → Séance → Étudiant",
                    "Notes : Filière → Semestre → UE → Étudiant",
                    "Enseignants : Filière → Cours enseignés → Enseignant",
                ],
                "etat": "A PLANIFIER",
            },
            {
                "nom": "2. Sécurité des fichiers (optionnel)",
                "etapes": [
                    "Purge des fichiers orphelins jamais validés",
                    "Vérification d'intégrité fichier ↔ base de données",
                    "Procédure de sauvegarde du dossier public/",
                ],
                "etat": "OPTIONNEL",
            },
        ],
    },
]


def ajouter_entete(document, titre):
    doc_entete = document.sections[0].header
    p = doc_entete.paragraphs[0]
    p.text = "EasyEcole — Découpage du projet en processus métier"
    p.runs[0].font.color.rgb = GRIS
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.name = "Calibri"


def ajouter_pied(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EasyEcole — Processus métier par pôle")
    run.font.size = Pt(9)
    run.font.color.rgb = GRIS


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Page de titre
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.add_run("EASYECOLE")
    run.font.size = Pt(34)
    run.bold = True
    run.font.color.rgb = BLEU

    sous = doc.add_paragraph()
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sous.add_run("Découpage du projet en processus métier")
    run.font.size = Pt(18)
    run.font.color.rgb = GRIS

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "Ce document organise le système de gestion académique et ERP comptable "
        "EasyEcole en processus métier, regroupés par pôle. Chaque processus décrit "
        "les étapes métier, du déclencheur jusqu'à la clôture, et indique son état "
        "d'avancement dans le projet."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # Tableau de synthèse
    doc.add_heading("Synthèse des processus", level=1)
    tableau = doc.add_table(rows=1, cols=3)
    tableau.style = "Light Grid Accent 1"
    tableau.alignment = WD_TABLE_ALIGNMENT.CENTER
    entetes = tableau.rows[0].cells
    for i, txt in enumerate(["Pôle", "Processus", "État"]):
        entetes[i].text = txt
        for par in entetes[i].paragraphs:
            for r in par.runs:
                r.bold = True

    for pole in PÔLES:
        for proc in pole["processus"]:
            row = tableau.add_row().cells
            row[0].text = pole["titre"].replace("PÔLE ", "").replace(" —", "")
            row[1].text = proc["nom"].split(". ", 1)[-1]
            row[2].text = proc["etat"]
            for par in row[2].paragraphs:
                for r in par.runs:
                    if "PLANIFIER" in proc["etat"] or "COMPLÉTER" in proc["etat"]:
                        r.font.color.rgb = RGBColor(0xB2, 0x6A, 0x00)
                    elif proc["etat"] == "LIVRÉ":
                        r.font.color.rgb = VERT
                    else:
                        r.font.color.rgb = RGBColor(0x9A, 0x1B, 0x1B)

    # Détail par pôle
    doc.add_page_break()
    for pole in PÔLES:
        doc.add_heading(pole["titre"], level=1)
        for proc in pole["processus"]:
            nom, etat = proc["nom"], proc["etat"]
            doc.add_heading(f"{nom}  —  Statut : {etat}", level=2)
            doc.add_paragraph("Étapes du processus :")
            for etape in proc["etapes"]:
                doc.add_paragraph(etape, style="List Number")
            doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("Prochaines priorités", level=1)
    priorités = [
        "Lot 6 — Clôture de semestre : gel des notes, validation UE, dettes académiques, passation.",
        "Programme ERP SYSCOHADA (S0-S4) : états financiers conformes OHADA (bilan, CR, TAFIRE, ETATC).",
        "Programme RH (R1-R5) : paie des prestataires, HAO depuis les pointages, charges sociales Togo.",
        "E-Learning arborescent (EL1-EL3) : rattachement hiérarchique Année → Parcours → Filière → Niveau.",
        "Arborescence généralisée (Lot 7) : effectifs, présences, notes, enseignants.",
        "Sécurité des fichiers : purge des orphelins, intégrité fichier ↔ base, sauvegarde public/.",
    ]
    for p in priorités:
        doc.add_paragraph(p, style="List Bullet")

    ajouter_entete(doc, titre)
    ajouter_pied(doc)

    out = os.path.join("D:\\EasyEcole", "PROCESSUS-METIER-EasyEcole.docx")
    doc.save(out)
    print(f"Document généré : {out}")


if __name__ == "__main__":
    main()
